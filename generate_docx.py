import sys
import os
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing it now...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    except Exception as e:
        print(f"Error installing python-docx: {e}")
        print("Please run: pip install python-docx")
        sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set custom cell padding (margins) in dxa (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    """Add clean grid borders to a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    # Top border
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '4')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), 'D3D3D3')
    tblBorders.append(top)
    
    # Bottom border
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '0')
    bottom.set(qn('w:color'), 'D3D3D3')
    tblBorders.append(bottom)
    
    # Inside horizontal borders
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:space'), '0')
    insideH.set(qn('w:color'), 'E0E0E0')
    tblBorders.append(insideH)

    # Left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '4')
    left.set(qn('w:color'), 'D3D3D3')
    tblBorders.append(left)

    # Right border
    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '4')
    right.set(qn('w:color'), 'D3D3D3')
    tblBorders.append(right)

    tblPr.append(tblBorders)

def generate_document():
    doc = Document()

    # Set page margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    def add_chapter(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 51, 102) # Dark Blue
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_heading2(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 100, 0) # Goldish/Brown
        return p

    def add_heading3(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.bold = True
        p.add_run(text)
        return p

    # ------------------ CHAPTER 6 ------------------
    add_chapter("CHAPTER SIX: SYSTEM IMPLEMENTATION")
    
    doc.add_paragraph(
        "System implementation is the phase where the theoretical design of the system is converted into a fully "
        "functioning software solution. This chapter details how the various modules, databases, and interfaces of "
        "the Panda Ngazi (TUK Bursary & Scholarship Portal) interact with each other to form a cohesive system. "
        "It discusses the technical environment, coding tools, debugging techniques, comprehensive test cases, "
        "change-over methodologies, and database testing schema."
    )

    add_heading2("6.1 Coding/Environment/Debugging Techniques")
    doc.add_paragraph(
        "The development and deployment of the Panda Ngazi portal utilized modern software engineering tools "
        "and environments to guarantee system stability, scalability, and security."
    )

    add_heading3("6.1.1 Coding Tools")
    doc.add_paragraph(
        "The core coding tools used for developing the application include:"
    )
    add_bullet(" Chosen for its rapid development capabilities, built-in security features (such as protection against CSRF, SQL Injection, and XSS), and the Model-Template-View (MTV) architectural pattern.", "Django Web Framework (Version 4.2+):")
    add_bullet(" The primary backend scripting language, known for readability, maintainability, and robust libraries.", "Python (Version 3.10+):")
    add_bullet(" Used as the local relational database management system (RDBMS) for storing user profiles, applications, feedback messages, and payment slips due to its lightweight nature and zero-configuration requirement.", "SQLite Database:")
    add_bullet(" Used to construct the front-end structure and a premium user interface with a custom typography grid, custom color variables (TUK colors: Gold and Blue), and mobile responsiveness.", "HTML5 & CSS3 (Vanilla CSS):")
    add_bullet(" Used for micro-animations, client-side notifications, confirmation boxes, and auto-dismissing toast messages.", "JavaScript (ES6):")
    add_bullet(" Used as the primary Integrated Development Environment (IDE) with extensions for Django, Python linting, and Git integration.", "Visual Studio Code (VS Code):")

    add_heading3("6.1.2 Environment")
    doc.add_paragraph(
        "The system execution environment was configured as follows:"
    )
    add_bullet(" Windows 10/11 for local development and coding, and Linux/Ubuntu for the proposed staging/deployment production server.", "Operating System:")
    add_bullet(" Used to isolate the application's dependencies and prevent version conflicts with system-wide Python libraries.", "Python Virtual Environment (venv):")
    add_bullet(" Used to host the application locally at http://127.0.0.1:8000/ during the building and debugging phases.", "Django Development Server:")
    add_bullet(" Used to install and manage third-party Python packages such as django-widget-tweaks.", "Pip Package Manager:")

    add_heading3("6.1.3 Debugging Tools")
    doc.add_paragraph(
        "To ensure code correctness and troubleshoot runtime errors, the following debugging tools were utilized:"
    )
    add_bullet(" Provided detailed traceback outputs in the browser during development, pinpointing lines of code causing execution failures.", "Django Interactive Error Page (DEBUG = True):")
    add_bullet(" Provided a command-line interface to interact directly with the database models to verify data creation and logic.", "Django Management Shell (python manage.py shell):")
    add_bullet(" Used for styling inspection, CSS grid adjustments, tracking HTTP requests (Network tab), and debugging JavaScript events.", "Web Browser Developer Tools (Chrome DevTools):")
    add_bullet(" Used to set breakpoints in views.py and step through functions to inspect variable values.", "Python Native Debugger (pdb / VS Code Debugger):")

    add_heading2("6.2 Program Listing")
    doc.add_paragraph(
        "The application is organized into a modular directory structure using Django's application design. "
        "The key modules and their file locations are listed below:"
    )
    add_bullet(" Contains project configurations, database definitions, middleware, and email validation settings.", "panda_ngazi/ (Project Config): settings.py, urls.py")
    add_bullet(" Defines custom User and StudentProfile models, and the email-based backend auth backend.", "accounts/ (User Authentication): models.py, backends.py")
    add_bullet(" Contains forms and views for submitting application details, uploading documents, downloading templates, and performing reviews.", "applications/ (Bursary Application): models.py, views.py, forms.py")
    add_bullet(" Implements file uploading fields for students to upload proof of bank receipts.", "payments/ (Disbursement Receipt): models.py, views.py")
    add_bullet(" Stores feedback details and logs review comments and status outcomes from admins.", "feedback/ (Messages & Queries): models.py, views.py")
    add_bullet(" Performs mathematical aggregates of requested amounts, total disbursements, and counts for application pipelines.", "reports/ (Analytics Dashboard): views.py")

    add_heading2("6.3 System/Program Testing")
    doc.add_paragraph(
        "Testing represents the comparison between expected behavior and actual system behavior. For the Panda Ngazi portal, "
        "Unit Testing and Integration Testing were carried out to evaluate each form, view validation, and database trigger. "
        "Below are the details of the key test cases executed during system evaluation:"
    )

    def add_test_case_table(summary, prereqs, procedures, test_data, expected, actual, status):
        table = doc.add_table(rows=7, cols=2)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_borders(table)
        
        headers = [
            ("Test Case Summary", summary),
            ("Prerequisites", prereqs),
            ("Test Procedures", procedures),
            ("Test Data", test_data),
            ("Expected Result", expected),
            ("Actual Results", actual),
            ("Status", status)
        ]
        
        for idx, (label, val) in enumerate(headers):
            row = table.rows[idx]
            
            # Label Cell
            cell_lbl = row.cells[0]
            cell_lbl.width = Inches(1.8)
            p_lbl = cell_lbl.paragraphs[0]
            p_lbl.paragraph_format.line_spacing = 1.15
            p_lbl.paragraph_format.space_after = Pt(2)
            run_lbl = p_lbl.add_run(label)
            run_lbl.bold = True
            
            # Value Cell
            cell_val = row.cells[1]
            cell_val.width = Inches(4.7)
            p_val = cell_val.paragraphs[0]
            p_val.paragraph_format.line_spacing = 1.15
            p_val.paragraph_format.space_after = Pt(2)
            
            if label == "Status":
                run_val = p_val.add_run(val)
                run_val.bold = True
                run_val.font.color.rgb = RGBColor(0, 128, 0) if "Pass" in val else RGBColor(255, 0, 0)
            else:
                p_val.add_run(val)
                
            set_cell_margins(cell_lbl)
            set_cell_margins(cell_val)
            
        doc.add_paragraph() # Spacer after table

    # Table 1
    add_heading3("Test Case 1: Student Registration and Domain Enforcement")
    add_test_case_table(
        "Verify student registration restricts signups to TUK student email domains.",
        "None",
        "1. Navigate to the signup page.\n2. Fill in the name, phone number, and password.\n3. Enter email 'johndoe@gmail.com' and click Register.\n4. Re-enter the same details but with email 'johndoe@students.tuk.ac.ke'.",
        "User-generated signup form inputs (TUK student vs. Gmail).",
        "Registration fails for the non-TUK domain (showing a validation error) and succeeds for the students.tuk.ac.ke domain.",
        "Registration blocked for Gmail; successfully registered student profile for TUK email.",
        "Pass"
    )

    # Table 2
    add_heading3("Test Case 2: Submitting a Bursary Application")
    add_test_case_table(
        "Verify that a registered student can fill out and submit the application form.",
        "Student must be registered and signed in.",
        "1. Click on the 'Apply' tab in navigation.\n2. Enter requested amount, vulnerability details, and family background.\n3. Upload supporting files (PDF fee statement).\n4. Click Submit.",
        "Hardship statements, amount: KES 15,000, fee statement upload.",
        "Application is successfully written to database, and user is redirected to the status page showing 'Pending Review'.",
        "Application record created in database; redirect executed and status updated.",
        "Pass"
    )

    # Table 3
    add_heading3("Test Case 3: Administrator Application Review")
    add_test_case_table(
        "Verify that administrators can filter, review, and update application statuses.",
        "Admin user must be logged in; at least one student application must be submitted.",
        "1. Navigate to the 'Applications' tab on the admin dashboard.\n2. Click on a pending application.\n3. Change status from 'Pending Review' to 'Approved'.\n4. Add feedback notes and submit.",
        "Status value: approved, Admin note: 'Approved for allocation of KES 12,000'.",
        "The application status is updated in the database, and the student dashboard immediately reflects the approved status.",
        "Status updated to 'Approved'; student view correctly displays step progression.",
        "Pass"
    )

    # Table 4
    add_heading3("Test Case 4: Payment Slip/Receipt Upload")
    add_test_case_table(
        "Verify that students can upload proof of receipt to confirm they deposited the bursary cheque.",
        "Student application must be marked as 'Disbursed' or 'Approved'.",
        "1. Go to 'Payments' tab.\n2. Select the approved application.\n3. Enter bank name and amount received.\n4. Upload the scan of the bank deposit receipt.\n5. Click Upload.",
        "Bank Name: 'Cooperative Bank', Amount: KES 12,000, file receipt image.",
        "Receipt entry is saved as unverified in the database, waiting for admin approval.",
        "Payment instance successfully linked to application and student.",
        "Pass"
    )

    # Table 5
    add_heading3("Test Case 5: Feedback Messaging and Notification")
    add_test_case_table(
        "Verify that students receive specific messages sent by admins about application corrections.",
        "Administrator and student accounts active; application under review.",
        "1. Admin navigates to student's application.\n2. Writes feedback requesting a clearer upload of the family fee statement.\n3. Saves feedback.\n4. Student logs in and navigates to the 'Messages' tab.",
        "Message: 'Please re-upload a clear copy of your official fee statement.'",
        "The message is saved in the database with status is_read = False and displays prominently on the student's dashboard.",
        "Feedback entry created; student message dashboard displays the notification.",
        "Pass"
    )

    add_heading2("6.4 Proposed Change-Over Techniques")
    doc.add_paragraph(
        "To transition the Technical University of Kenya from the legacy, paper-based manual bursary application "
        "method to the new online portal, a Pilot Change-over combined with a short window of Parallel Change-over is proposed:"
    )
    add_bullet(" The system will first be launched exclusively for students within the School of Computing and Information Technology (SCIT). Testing the system with a single school allows the development team to monitor server loads, user experience difficulties, and database query performances under actual use constraints.", "1) Pilot Change-over:")
    add_bullet(" During the first semester of pilot roll-out, the bursary committee will keep physical paper application dropboxes active alongside the portal. This guarantees a fallback mechanism in case of unexpected hosting outages or data loss, ensuring that no student is locked out of financial support during system adoption.", "2) Parallel Change-over:")
    add_bullet(" Once the pilot is successful and the system has been refined based on feedback, the university will transition to a complete online-only submission format for all faculties, rendering the paper system obsolete.", "3) Direct Cut-over:")

    add_heading2("6.5 Test Data")
    doc.add_paragraph(
        "Test data was seeded in the SQLite database (db.sqlite3) using Django fixtures and manual entries. "
        "The relational structure tested includes:"
    )
    add_bullet(" Seeded with 5 student accounts (with emails ending in @students.tuk.ac.ke) and 2 administrator accounts (ending in @tuk.ac.ke).", "auth_user:")
    add_bullet(" Populated with details representing different years of study (Year 1 to Year 5), mock course names (e.g., 'Bachelor of Technology in Information Technology'), and varying vulnerability statuses (e.g., 'Orphan', 'Single parent background').", "accounts_studentprofile:")
    add_bullet(" Seeded with test entries representing various stages of the review lifecycle (pending, qualified, approved, and rejected) to test admin dashboard visualizations and filters.", "applications_application:")

    add_heading2("6.6 Sample Run - Output")
    doc.add_paragraph(
        "The system was compiled and hosted on a local server. During the sample execution run:"
    )
    add_bullet(" Database tables were checked using SQLite database tools. Relations were correctly linked through foreign keys (such as student_id in the application pointing to the custom user model).", "1) Database Inspection:")
    add_bullet(" Navigations successfully loaded dynamic content. Submitting a form generated instant toast message banners at the top of the browser window.", "2) Page Flow Verification:")
    add_bullet(" Aggregated reports correctly counted application totals and financial sums.", "3) Admin Portal Evaluation:")


    # ------------------ CHAPTER 7 ------------------
    doc.add_page_break()
    add_chapter("CHAPTER SEVEN: USER MANUAL - DOCUMENTATION")
    
    add_heading2("7.1 Installation Environment")
    doc.add_paragraph(
        "The Panda Ngazi application is built to run on standard operating systems. In a local development environment, "
        "it can be run on any PC with Windows, macOS, or Linux, using the lightweight built-in Python web server. "
        "In a production hosting environment, the system is designed to run on a Linux Web Server (e.g., Ubuntu Server 22.04 LTS) "
        "using a WSGI server like Gunicorn managed under an Nginx reverse proxy server. Media files (such as students' PDF certificates) "
        "are stored in the server's local storage directory under the /media/ URL route."
    )

    add_heading2("7.2 Installation Requirements")
    add_heading3("7.2.1 Hardware Requirements")
    add_bullet(" Intel Core i3 / AMD Ryzen 3 or higher (Minimum 2.0 GHz clock speed).", "Processor:")
    add_bullet(" Minimum 4 GB RAM (8 GB recommended for multitasking during development).", "Memory:")
    add_bullet(" At least 500 MB of available disk space for system files and dependencies (additional space required based on database growth and file uploads).", "Storage Space:")

    add_heading3("7.2.2 Software Requirements")
    add_bullet(" Windows 10/11, macOS Monterey+, or Linux Ubuntu 20.04+.", "Operating System:")
    add_bullet(" Python Version 3.10 or higher.", "Runtime Environment:")
    add_bullet(" SQLite 3 (included with Python).", "Database:")
    add_bullet(" Any modern CSS3/HTML5 compliant browser (Google Chrome, Mozilla Firefox, Apple Safari, Microsoft Edge).", "Web Browser:")

    add_heading2("7.3 Installation Procedures")
    doc.add_paragraph(
        "To download and set up the project files on your local machine, follow these instructions:"
    )
    
    procedures_list = [
        ("1. Download Codebase", "Extract the project ZIP folder or clone the repository to your desktop using:\ngit clone https://github.com/yourusername/panda-ngazi.git\ncd panda-ngazi"),
        ("2. Initialize Virtual Environment", "Open terminal/PowerShell inside the project root folder and execute:\npython -m venv venv"),
        ("3. Activate Virtual Environment", "Windows (PowerShell):\n.\\venv\\Scripts\\Activate.ps1\n\nMac/Linux (Terminal):\nsource venv/bin/activate"),
        ("4. Install Required Libraries", "Install packages listed in the requirements file:\npip install -r requirements.txt"),
        ("5. Configure Database Schema", "Run database migrations to generate database tables:\npython manage.py makemigrations\npython manage.py migrate"),
        ("6. Create Administration User", "Generate the master admin account to access the control panel:\npython manage.py createsuperuser\n(Follow the prompts to enter an admin email ending in @tuk.ac.ke and set a password)."),
        ("7. Start Development Server", "Run the server using:\npython manage.py runserver"),
        ("8. Access App", "Open http://127.0.0.1:8000/ in a web browser.")
    ]
    
    for step_title, step_desc in procedures_list:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        run_title = p.add_run(f"{step_title}: ")
        run_title.bold = True
        p.add_run(step_desc)

    add_heading2("7.4 User Instructions")
    
    add_heading3("7.4.1 User Sign Up for an Account")
    doc.add_paragraph(
        "1. Open the portal homepage and click the Register button in the top right.\n"
        "2. Fill in the registration form. Crucial: If you are a student, your email must end with the domain @students.tuk.ac.ke. "
        "If you are a staff member or administrator, your email must end with @tuk.ac.ke.\n"
        "3. Click Sign Up. Upon success, you will be redirected to the sign-in screen.\n"
        "4. On first login, students will be prompted to fill out their Student Profile including their Student Registration Number "
        "(e.g., SCII/XXXXX/202X), Course Name, Year of Study, and basic Family Background description."
    )

    add_heading3("7.4.2 Sign In to Account")
    doc.add_paragraph(
        "1. Click Sign In on the navigation bar.\n"
        "2. Enter your registered email address and password.\n"
        "3. Click the Sign In button.\n"
        "4. The system automatically checks your role: Students are redirected to the student dashboard, "
        "and Administrators are redirected to the administrative reports page."
    )

    add_heading3("7.4.3 Home/Dashboard Screens")
    doc.add_paragraph(
        "Student Home: Displays a progress bar representing the current application lifecycle stage (Pending Review -> Qualified -> Approved -> Disbursed). "
        "Includes shortcuts to update application details, view feedback messages, and upload payment receipts.\n\n"
        "Admin Dashboard: Displays statistics panels including: Total Applications Submitted, Pending Review Count, "
        "Approved vs. Disbursed Allocation Totals, and Active Student Registrations. It displays a list of the 10 most recent submissions. "
        "Admins can click on any application to see student details, inspect uploaded certificates, input internal notes, and write feedback."
    )

    add_heading3("7.4.4 How to Reset Password")
    doc.add_paragraph(
        "1. On the Login screen, click the Forgot Password? link below the input fields.\n"
        "2. Enter the email address associated with your account.\n"
        "3. Click Reset Password.\n"
        "4. In development: The password reset link will be printed directly to the system console terminal.\n"
        "5. In production: A reset link will be sent to the student's email inbox, allowing them to click it and input a new password."
    )

    add_heading2("7.5 System Conversion Method")
    doc.add_paragraph(
        "System conversion will adopt a schema mapping procedure. The existing database structure uses SQLite. To migrate old records:\n"
        "1. Bursary office staff will export current Excel sheets of past bursary recipients into standard CSV files.\n"
        "2. A custom Django command (python manage.py import_students csv_file.csv) will read student numbers, verify domains, "
        "and programmatically insert records into accounts_user and accounts_studentprofile tables.\n"
        "3. This automated script minimizes manual keying errors and speeds up the deployment cycle."
    )

    add_heading2("7.6 User Training")
    doc.add_paragraph(
        "To ensure high system adoption rates:\n"
        "- Student Training: A brief user guide page with clear instructions will be accessible on the portal without requiring sign-in. "
        "Dynamic input validation guidelines are embedded directly inside the forms to assist students as they fill in fields.\n"
        "- Admin Training: A half-day training workshop will be conducted for the TUK Bursary Committee members. The training covers "
        "using application filters, uploading template blank forms, entering feedback, and generating reports."
    )

    add_heading2("7.7 File Conversions")
    doc.add_paragraph(
        "- Supporting Certificates: The system automatically converts uploaded images (PNG, JPEG) and PDF files "
        "and stores them in structured folders under /media/applications/supporting/.\n"
        "- Template Forms: Blank applications uploaded by administrators are saved in /media/form_templates/ "
        "as downloadable files for students."
    )


    # ------------------ CHAPTER 8 ------------------
    doc.add_page_break()
    add_chapter("CHAPTER EIGHT: LIMITATIONS, CHALLENGES, CONCLUSIONS AND RECOMMENDATIONS")
    
    add_heading2("8.1 Limitations")
    doc.add_paragraph(
        "While the Panda Ngazi portal provides a more modern alternative to the old manual system, it faces the following constraints:"
    )
    add_bullet(" The application requires active internet access. Students in remote locations with poor network coverage may face challenges when attempting to upload large PDF files (e.g., high-resolution fee statements).", "1. Internet Connectivity Dependency:")
    add_bullet(" Although the system automates application organization, filtering, and communication, the actual evaluation of the authenticity of student documents (such as verified death certificates of guardians or chief letters) must still be carried out manually by the committee.", "2. Manual Verification of Documents:")
    add_bullet(" Using SQLite as a local file-based database limits application write capacity if hundreds of students attempt to submit applications at the same time during peak deadline hours.", "3. Stateless Local Session Storage:")

    add_heading2("8.2 Challenges Faced")
    doc.add_paragraph(
        "During the system development phase, the following challenges were encountered:"
    )
    add_bullet(" Designing a responsive HTML/CSS visual timeline tracker that dynamically reflects the database application status step-by-step required complex template logic.", "1. Dynamic Status Tracker Implementation:")
    add_bullet(" Enforcing domain validations on the registration form to differentiate students (@students.tuk.ac.ke) from staff (@tuk.ac.ke) required override validations in Django forms, which was complex to troubleshoot initially.", "2. Strict Email Domain Filters:")
    add_bullet(" Configuring the system to handle large file sizes without crashing the local development server required fine-tuning configurations and setting limits on file upload sizes in the frontend form inputs.", "3. Media Upload File Limitations:")

    add_heading2("8.3 Degree of Success")
    doc.add_paragraph(
        "The project achieved a high degree of success based on its objectives:\n"
        "- 100% Online Registration and Form Submission: Students can complete profiles and submit all necessary files online.\n"
        "- Clear Administrative Pipeline: The bursary committee can search, filter by status, write feedback, and process applications without paper files.\n"
        "- Automated Tracking: Students can monitor the status of their application, reducing physical inquiry visits to the student affairs offices."
    )

    add_heading2("8.4 Learning Experience")
    doc.add_paragraph(
        "Developing the Panda Ngazi portal provided a valuable learning experience:\n"
        "- Framework Architecture: Gained deep knowledge of Django's MVC/MTV design pattern, modular app organization, routing mechanisms, and templating engines.\n"
        "- Database Management: Learned table modeling, one-to-one and foreign key mappings, and optimized ORM queries.\n"
        "- UI/UX Design: Acquired skills in building clean user interfaces using vanilla CSS custom properties, fluid layouts, grid designs, and accessible forms."
    )

    add_heading2("8.5 Recommendations")
    doc.add_paragraph(
        "For future versions of the system, the following updates are recommended:"
    )
    add_bullet(" Integrate the portal directly with the main university registrar database. This would enable automatic verification of student admission status, GPA, and outstanding fee balance, eliminating document forgery.", "1. API Integration with TUK Student Portal:")
    add_bullet(" Integrate SMS gateways (such as Africa's Talking) to automatically send mobile text notifications to students when their bursary applications are approved or when corrections are needed.", "2. SMS API Integration:")
    add_bullet(" Implement libraries (such as ReportLab or WeasyPrint) allowing admins to export aggregated tables and charts as PDF reports for university board presentations.", "3. PDF Report Exporter:")
    add_bullet(" Migrate from SQLite to a robust PostgreSQL database server to support high concurrency during peak application periods.", "4. Database Upgrade:")

    add_heading2("8.6 Conclusion")
    doc.add_paragraph(
        "The Panda Ngazi (TUK Bursary & Scholarship Portal) successfully addresses the inefficiencies, lack of transparency, "
        "and delays associated with the manual application system. By digitizing student profile creation, bursary submissions, "
        "admin reviews, and communication, the system creates a faster, more accountable flow for financial aid distribution at the "
        "Technical University of Kenya. Implementing the recommendations detailed above will further secure, optimize, and scale the portal, "
        "helping the university support needy students effectively."
    )

    # Save document
    output_filename = "PANDA_NGAZI_DOCUMENTATION_CH_6-8.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}!")

if __name__ == "__main__":
    generate_document()
