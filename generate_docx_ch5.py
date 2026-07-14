import sys
import os
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("python-docx not found. Installing it now...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        import docx
    except Exception as e:
        print(f"Error installing python-docx: {e}")
        print("Please run: pip install python-docx")
        sys.exit(1)

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Set custom cell padding (margins) in dxa (1/20 of a point)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_table_borders(table):
    """Add clean grid borders to a table."""
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    
    for side in ['top', 'bottom', 'left', 'right', 'insideH']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'D3D3D3' if side != 'insideH' else 'E0E0E0')
        tblBorders.append(border)
        
    tblPr.append(tblBorders)

def generate_ch5_document():
    doc = Document()

    # Set page margins to 1 inch
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base style setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 1.5
    style_normal.paragraph_format.space_after = Pt(6)

    def add_chapter(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 51, 102) # TUK Navy
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_heading2(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(128, 100, 0) # Goldish/Brown
        return p

    def add_heading3(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_prefix = p.add_run(bold_prefix)
            run_prefix.bold = True
        p.add_run(text)
        return p

    def add_ascii_block(text):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.0
        
        # Use a light background highlight for block
        pPr = p._p.get_or_add_pPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), 'F4F4F6')
        pPr.append(shading)
        
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9.5)

    def add_data_dict_table(headers, rows_data):
        table = doc.add_table(rows=len(rows_data) + 1, cols=6)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_table_borders(table)
        
        # Header Row
        hdr_row = table.rows[0]
        widths = [Inches(1.2), Inches(0.8), Inches(0.8), Inches(0.6), Inches(0.6), Inches(2.5)]
        
        for idx, name in enumerate(headers):
            cell = hdr_row.cells[idx]
            cell.width = widths[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name)
            run.bold = True
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            
        # Data Rows
        for row_idx, row_data in enumerate(rows_data):
            row = table.rows[row_idx + 1]
            for col_idx, text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.width = widths[col_idx]
                p = cell.paragraphs[0]
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                p.add_run(str(text))
                set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                
        doc.add_paragraph() # Spacer

    # ------------------ TITLE ------------------
    add_chapter("CHAPTER FIVE: SYSTEM DESIGN")
    
    # 5.1
    add_heading2("5.1 Introduction to System Design and Nature of the System")
    doc.add_paragraph(
        "System design is the process of defining the architecture, components, modules, interfaces, and data for a "
        "system to satisfy specified requirements. This chapter details the logical and physical design of the Panda "
        "Ngazi (TUK Bursary & Scholarship Portal)."
    )
    doc.add_paragraph(
        "Unlike mobile-based applications that run on client operating systems, the Panda Ngazi portal is a server-side "
        "web application built on the Model-Template-View (MTV) architectural pattern of the Django framework. The nature of "
        "this system is an interactive, double-role portal (serving Students and Administrators) that operates on relational "
        "database management paradigms to organize application files, track qualifications, manage tuition disbursements, "
        "and display reports."
    )

    # 5.2
    add_heading2("5.2 Design Objectives")
    doc.add_paragraph(
        "To ensure system success, the design of the portal was guided by the following objectives:"
    )
    add_bullet(" The user interfaces (UI) must be clean and simple, enabling students of any technological background to easily register, download forms, upload documents, and track status.", "1. User-Friendliness:")
    add_bullet(" The payment module is designed to collect official receipts from the school finance office instead of generic bank slips. This ensures that the disbursed money is credited directly to the student’s fee account, reducing cash fraud.", "2. Accountability and Fraud Mitigation:")
    add_bullet(" The feedback module must provide a real-time progress tracker (stepper) showing exactly where the student’s application is: Pending Review -> Qualified -> Approved -> Disbursed.", "3. Real-time Status Transparency:")
    add_bullet(" The system must provide clean charts and downloadable data reports (CSV format) for administrators to present to the university bursary board.", "4. Data Monitoring:")

    # 5.3
    add_heading2("5.3 Program Design Tools")
    doc.add_paragraph(
        "The logic and structure of the Panda Ngazi system were designed using standard software engineering modeling tools, as described below."
    )

    add_heading3("5.3.1 Flow Chart")
    doc.add_paragraph(
        "The flowchart represents the sequential steps a user takes from entering the landing page to completing their tasks. The flowchart below visualizes the two distinct user paths:"
    )
    
    flowchart_ascii = (
        "                  +-------------------+\n"
        "                  |       Start       |\n"
        "                  +---------+---------+\n"
        "                            |\n"
        "                            v\n"
        "                  +---------+---------+\n"
        "                  |    Is Logged In?  |<------------------------+\n"
        "                  +----+---------+----+                         |\n"
        "                       |         |                              |\n"
        "                   No  |         | Yes                          |\n"
        "                       v         v                              |\n"
        "            +----------+---+   +-+----------------------------+ |\n"
        "            | Choose Role  |   | Identify Role                | |\n"
        "            +---+------+---+   +-+--------------+-------------+ |\n"
        "                |      |         |              |               |\n"
        "        Student |      | Admin   | Student      | Admin         |\n"
        "                v      v         v              v               |\n"
        "            +---+--+ +-+----+  +-+-----------+  +-------------+ |\n"
        "            |SignUp| |SignUp|  |Dashboard    |  |Dashboard    | |\n"
        "            +---+--+ +-+----+  +-+-----+-----+  +------+------+ |\n"
        "                |      |               |               |        |\n"
        "                +------+               |               |        |\n"
        "                       |               v               |        |\n"
        "                       |       +-------+-------+       |        |\n"
        "                       |       |Download Form  |       |        |\n"
        "                       |       +-------+-------+       |        |\n"
        "                       |               |               |        |\n"
        "                       |       +-------+-------+       |        |\n"
        "                       |       |Upload Filled  |       |        |\n"
        "                       |       |Form & ID      |       |        |\n"
        "                       |       +-------+-------+       |        |\n"
        "                       |               |               |        |\n"
        "                       |               v               v        |\n"
        "                       |       +-------+-------+ +-----+-----+  |\n"
        "                       |       |Track Status   | |Review     |  |\n"
        "                       |       |Stepper        | |Application|  |\n"
        "                       |       +-------+-------+ +-----+-----+  |\n"
        "                       |               |               |        |\n"
        "                       |               v               v        |\n"
        "                       |       +-------+-------+ +-----+-----+  |\n"
        "                       |       |Upload School  | |Verify     |  |\n"
        "                       |       |Fee Receipt    | |Receipts   |  |\n"
        "                       |       +-------+-------+ +-----+-----+  |\n"
        "                       |               |               |        |\n"
        "                       |               |               v        |\n"
        "                       |               |         +-----+-----+  |\n"
        "                       |               |         |Export CSV |  |\n"
        "                       |               |         |Reports    |  |\n"
        "                       |               |         +-----+-----+  |\n"
        "                       |               v               |        |\n"
        "                       +-------------->+---------------+--------+\n"
        "                                       |\n"
        "                                       v\n"
        "                               +-------+-------+\n"
        "                               |     Logout    |\n"
        "                               +-------+-------+\n"
        "                                       |\n"
        "                                       v\n"
        "                               +-------+-------+\n"
        "                               |     Stop      |\n"
        "                               +---------------+\n"
    )
    add_ascii_block(flowchart_ascii)

    add_heading3("5.3.2 Use Case Diagram")
    doc.add_paragraph(
        "The Use Case diagram outlines how the two key actors (Student and Administrator) interact with the system's boundary features."
    )
    
    usecase_ascii = (
        "                      +---------------------------------------+\n"
        "                      |          PANDA NGAZI SYSTEM           |\n"
        "                      |                                       |\n"
        "                      |   +-------------------------------+   |\n"
        "                      |   |        Register Account       |   |\n"
        "                      |   +---------------+---------------+   |\n"
        "                      |                   ^                   |\n"
        "                      |                   |                   |\n"
        "      +---------+     |   +---------------+---------------+   |     +---------+\n"
        "      |         |-----+---|      Download Blank Form      |---|----|         |\n"
        "      |         |     |   +-------------------------------+   |    |         |\n"
        "      |         |-----+---|       Submit Application      |   |    |         |\n"
        "      | Student |     |   +-------------------------------+   |    |  Admin  |\n"
        "      |  Actor  |-----+---|      Track Status Stepper     |   |    |  Actor  |\n"
        "      |         |     |   +-------------------------------+   |    |         |\n"
        "      |         |-----+---|      View Feedback Notes      |---|----|         |\n"
        "      |         |     |   +-------------------------------+   |    |         |\n"
        "      |         |-----+---|     Upload School Receipt     |   |    |         |\n"
        "      +---------+     |   +-------------------------------+   |    +---------+\n"
        "                      |   |    Upload Blank Templates     |---|----|\n"
        "                      |   +-------------------------------+   |\n"
        "                      |   |     Review & Set Status       |---|----|\n"
        "                      |   +-------------------------------+   |\n"
        "                      |   |      Verify Fee Receipts      |---|----|\n"
        "                      |   +-------------------------------+   |\n"
        "                      |   |    Export CSV Data Reports    |---|----|\n"
        "                      |   +-------------------------------+   |\n"
        "                      +---------------------------------------+\n"
    )
    add_ascii_block(usecase_ascii)

    add_heading3("5.3.3 Data Flow Diagram (DFD)")
    doc.add_paragraph(
        "The Level-1 DFD models how data moves between processes, actors, and database stores:"
    )
    
    dfd_ascii = (
        "+---------+         Login Details         +---------+        User Credentials\n"
        "|  User   |==============================>| 1.0     |<============================+\n"
        "| (Actor) |<------------------------------| Auth    |--------------------------+  |\n"
        "+---------+         Auth Redirect         +---------+                          |  |\n"
        "    ||                                                                         v  v\n"
        "    ||                                                                    +-----------+\n"
        "    ||                                                                    |   D1:     |\n"
        "    ||  Filled Form & ID docs                                             | User DB   |\n"
        "    |+==================================================+                 +-----------+\n"
        "    ||                                                  |                      ^\n"
        "    ||                                                  v                      |\n"
        "    ||                                            +-----------+                |\n"
        "    ||                                            | 2.0       |================+\n"
        "    ||                                            | Submit    |    Application Details\n"
        "    ||                                            | App       |================+\n"
        "    ||                                            +-----------+                |\n"
        "    ||                                                                         v\n"
        "    ||                                                                    +-----------+\n"
        "    ||                                            +-----------+           |   D2:     |\n"
        "    ||               Review Comments / Status     | 3.0       |==========>| App DB    |\n"
        "    ||<-------------------------------------------| Application|          +-----------+\n"
        "    ||                                            | Review    |                ^\n"
        "    ||                                            +-----------+                |\n"
        "    ||                                                  ^                      |\n"
        "    ||                                                  |                      |\n"
        "    ||                                            +-----+-----+                |\n"
        "    ||                                            | 4.0       |================+\n"
        "    ||<===========================================| Feedback  |      Feedback Comments\n"
        "    ||               Feedback Log Notification    | Messaging |======>+\n"
        "    ||                                            +-----------+       |\n"
        "    ||                                                                v\n"
        "    ||                                                           +-----------+\n"
        "    ||                                                           |   D3:     |\n"
        "    ||                                            +-----------+  |Feedback DB|\n"
        "    ||               Finance Receipt Upload       | 5.0       |  +-----------+\n"
        "    |+===========================================>| Payment   |\n"
        "    ||                                            | Processing|======>\n"
        "    ||                                            +-----------+       |\n"
        "    ||                                                  ^             v\n"
        "    ||                                                  |        +-----------+\n"
        "    ||                                            +-----+-----+  |   D4:     |\n"
        "    ||<-------------------------------------------| 6.0       |<=|Payment DB |\n"
        "    ||               CSV Audit Spreadsheets       | Reports   |  +-----------+\n"
        "    +---------+                                   +-----------+\n"
    )
    add_ascii_block(dfd_ascii)

    # 5.4
    add_heading2("5.4 Logical Design")
    
    add_heading3("5.4.1 Logical Data Design (Entity Relationship Structure)")
    doc.add_paragraph(
        "The logical database layout maps relational structures in the backend. Users have a one-to-one relationship "
        "with a StudentProfile (if they are students). Users (Students) can submit multiple Applications over different "
        "academic cycles. Applications are reviewed by a User (Admin). Applications can contain multiple Feedback logs "
        "addressing the student. Applications are linked to Payments (when tuition accounts receive disbursements)."
    )
    
    er_ascii = (
        "  +------------------+             +------------------+\n"
        "  |      User        |1           1|  StudentProfile  |\n"
        "  | (auth_user)      |-------------| (studentprofile) |\n"
        "  +------------------+             +------------------+\n"
        "          |\n"
        "          |1\n"
        "          |\n"
        "          |*\n"
        "  +------------------+             +------------------+\n"
        "  |   Application    |1           *|     Feedback     |\n"
        "  |  (application)   |-------------|    (feedback)    |\n"
        "  +------------------+             +------------------+\n"
        "          |\n"
        "          |1\n"
        "          |\n"
        "          |*\n"
        "  +------------------+\n"
        "  |     Payment      |\n"
        "  |    (payment)     |\n"
        "  +------------------+\n"
    )
    add_ascii_block(er_ascii)

    add_heading3("5.4.2 Entity Life History")
    doc.add_paragraph(
        "The life history of a bursary application record progresses through specific state transitions:\n"
        "1. Creation: The record is initiated when a student submits their personal statement and files. Status is set to pending.\n"
        "2. Evaluation: An administrator evaluates the documents. If details are missing, feedback is sent. If parameters are met, the status transitions to qualified.\n"
        "3. Approval: The Bursary Board confirms financial allocation, updating the status to approved.\n"
        "4. Disbursement: The funds are sent to the student's TUK tuition account. The status changes to disbursed.\n"
        "5. Acknowledge/Verify: The student uploads the accounts office fee receipt, which is then verified by the admin.\n"
        "6. Termination (Alternative): If details are fraudulent or goals aren't met, status changes to rejected."
    )

    # 5.5
    add_heading2("5.5 Physical Design Description")
    doc.add_paragraph(
        "The physical design defines the database tables, fields, types, and user interfaces."
    )

    add_heading3("5.5.1 Data Dictionary")
    
    headers = ['Field Name', 'Data Type', 'Key Type', 'Size', 'Null?', 'Description']
    
    # User
    doc.add_paragraph("Table 1: User Table (accounts_user)")
    user_rows = [
        ['id', 'Integer', 'Primary', 'Auto', 'No', 'Unique user identifier'],
        ['username', 'Varchar', 'Unique', '150', 'No', 'URL-safe name derived from email'],
        ['email', 'Varchar', 'Unique', '254', 'No', 'Official TUK email address'],
        ['phone', 'Varchar', '-', '15', 'Yes', 'Contact phone number'],
        ['first_name', 'Varchar', '-', '150', 'No', 'User first name'],
        ['last_name', 'Varchar', '-', '150', 'No', 'User last name'],
        ['is_student', 'Boolean', '-', '1', 'No', 'Flag indicating student role'],
        ['is_admin', 'Boolean', '-', '1', 'No', 'Flag indicating admin role'],
        ['is_staff', 'Boolean', '-', '1', 'No', 'Flag permitting Django admin access'],
        ['is_active', 'Boolean', '-', '1', 'No', 'Account active status indicator'],
        ['password', 'Varchar', '-', '128', 'No', 'Hashed user password']
    ]
    add_data_dict_table(headers, user_rows)

    # StudentProfile
    doc.add_paragraph("Table 2: StudentProfile Table (accounts_studentprofile)")
    profile_rows = [
        ['user_id', 'Integer', 'Foreign', 'Auto', 'No', 'Links to User table (One-to-One)'],
        ['student_number', 'Varchar', 'Unique', '20', 'No', 'Student Admission Number'],
        ['course', 'Varchar', '-', '200', 'No', 'Enrolled Program of Study'],
        ['year_of_study', 'Integer', '-', '1', 'No', 'Current year (1 to 5)'],
        ['vulnerability_status', 'Text', '-', '-', 'Yes', 'Vulnerability details statement'],
        ['family_background', 'Text', '-', '-', 'Yes', 'Family financial background details']
    ]
    add_data_dict_table(headers, profile_rows)

    # ApplicationFormTemplate
    doc.add_paragraph("Table 3: ApplicationFormTemplate Table (applications_applicationformtemplate)")
    template_rows = [
        ['id', 'Integer', 'Primary', 'Auto', 'No', 'Unique Template ID'],
        ['title', 'Varchar', '-', '200', 'No', 'Title of the blank template'],
        ['description', 'Text', '-', '-', 'Yes', 'Helper text / download instructions'],
        ['form_file', 'Varchar (File)', '-', '100', 'No', 'Path to template file stored in media'],
        ['uploaded_at', 'DateTime', '-', 'Auto', 'No', 'Timestamp of file upload'],
        ['is_active', 'Boolean', '-', '1', 'No', 'True if this is the active downloadable form']
    ]
    add_data_dict_table(headers, template_rows)

    # Application
    doc.add_paragraph("Table 4: Application Table (applications_application)")
    app_rows = [
        ['id', 'Integer', 'Primary', 'Auto', 'No', 'Unique Application ID'],
        ['student_id', 'Integer', 'Foreign', 'Auto', 'No', 'Links to User table (Student)'],
        ['submitted_at', 'DateTime', '-', 'Auto', 'No', 'Initial submission timestamp'],
        ['updated_at', 'DateTime', '-', 'Auto', 'No', 'Last update timestamp'],
        ['status', 'Varchar', '-', '20', 'No', 'Cycle status (pending, qualified, approved, disbursed, rejected)'],
        ['vulnerability_details', 'Text', '-', '-', 'No', 'Statement of hardship'],
        ['family_background', 'Text', '-', '-', 'No', 'Description of family background'],
        ['amount_requested', 'Decimal', '-', '10, 2', 'Yes', 'Tuition amount requested (KES)'],
        ['supporting_documents', 'Varchar (File)', '-', '100', 'Yes', 'Path to supporting certificates'],
        ['filled_form', 'Varchar (File)', '-', '100', 'Yes', 'Path to completed form template'],
        ['reviewed_by_id', 'Integer', 'Foreign', 'Auto', 'Yes', 'Links to User (Admin reviewer)'],
        ['reviewed_at', 'DateTime', '-', 'Auto', 'Yes', 'Date reviewed by admin'],
        ['admin_notes', 'Text', '-', '-', 'Yes', 'Internal committee notes (private)'],
        ['feedback_to_student', 'Text', '-', '-', 'Yes', 'Message visible to the student']
    ]
    add_data_dict_table(headers, app_rows)

    # Feedback
    doc.add_paragraph("Table 5: Feedback Table (feedback_feedback)")
    feedback_rows = [
        ['id', 'Integer', 'Primary', 'Auto', 'No', 'Unique feedback ID'],
        ['application_id', 'Integer', 'Foreign', 'Auto', 'Yes', 'Links to corresponding application'],
        ['student_id', 'Integer', 'Foreign', 'Auto', 'No', 'Links to User table (Student)'],
        ['feedback_type', 'Varchar', '-', '30', 'No', 'Type of log (info, clarification, approval_note, rejection_reason)'],
        ['comments', 'Text', '-', '-', 'No', 'Feedback message body'],
        ['reviewed_by_id', 'Integer', 'Foreign', 'Auto', 'Yes', 'Admin sender ID'],
        ['reviewed_at', 'DateTime', '-', 'Auto', 'No', 'Date logged'],
        ['is_read', 'Boolean', '-', '1', 'No', 'Tracks if student viewed message']
    ]
    add_data_dict_table(headers, feedback_rows)

    # Payment
    doc.add_paragraph("Table 6: Payment Table (payments_payment)")
    payment_rows = [
        ['id', 'Integer', 'Primary', 'Auto', 'No', 'Unique receipt ID'],
        ['student_id', 'Integer', 'Foreign', 'Auto', 'No', 'Links to Student User'],
        ['application_id', 'Integer', 'Foreign', 'Auto', 'Yes', 'Links to Application'],
        ['amount', 'Decimal', '-', '10, 2', 'No', 'Bursary amount credited (KES)'],
        ['receipt', 'Varchar (File)', '-', '100', 'No', 'Path to receipt stored in media'],
        ['bank_name', 'Varchar', '-', '100', 'Yes', 'Finance office / Bank name'],
        ['uploaded_at', 'DateTime', '-', 'Auto', 'No', 'Date uploaded'],
        ['verified', 'Boolean', '-', '1', 'No', 'True if confirmed genuine by admin']
    ]
    add_data_dict_table(headers, payment_rows)

    # 5.5.2 to 5.5.5
    add_heading3("5.5.2 Database Design")
    doc.add_paragraph(
        "The database utilizes SQLite 3 as the RDBMS during local development. All tables are structured using Django’s "
        "Object-Relational Mapping (ORM) to handle integrity constraints. Relationships are enforced at the physical "
        "database layer using foreign keys linked to auth_user.id. Cascade deletion is configured on profile schemas. If "
        "a student account is removed, their profiles, applications, feedback, and payment receipts are automatically deleted "
        "to maintain database referential integrity."
    )

    add_heading3("5.5.3 Input Screen Design")
    doc.add_paragraph(
        "Input screen forms are styled using HTML5 inputs, styled drop-zones, and validation hints:\n"
        "1. Student Registration Form: Input fields for Name, email domain verification (checking for @students.tuk.ac.ke), phone, course, year, and password fields.\n"
        "2. Admin Registration Form: Fields for administrator name, official staff email (validated for @tuk.ac.ke), phone, and password inputs.\n"
        "3. Application Form Upload Interface: A Drag-and-Drop file upload area utilizing custom JavaScript to display selected file names. Validates file types (.pdf, .docx, .zip) before upload.\n"
        "4. Admin Review Form: Located inside the application details page. It provides a dropdown select box for status adjustments and text areas for messages and internal logs.\n"
        "5. Receipt Upload Interface: Simple inputs for the receipt amount, finance office name, and a file selector for the scanned image/PDF."
    )

    add_heading3("5.5.4 Output Screen Design")
    doc.add_paragraph(
        "Output displays are designed to structure tabular data and track applications:\n"
        "1. Student Tracker Stepper: A visual timeline stepper using CSS classes (current, done) and a dynamic percentage track background (0%, 33%, 66%, 100%) representing statuses.\n"
        "2. Analytics dashboard: Includes cards highlighting system statistics: total applications, pending queue count, and total disbursements (aggregated using SQL Sum query).\n"
        "3. CSV Reports: Generated by assembling database fields into tabular files. Administrators download these reports by clicking the Export Applications or Export Payments buttons."
    )

    add_heading3("5.5.5 Code Design")
    doc.add_paragraph(
        "The code architecture uses the Model-Template-View (MTV) pattern:\n"
        "- Model Layer (models.py): Represents SQL schemas as Python classes, keeping data logic isolated.\n"
        "- Template Layer (HTML/CSS): Renders the user interface. It utilizes Vanilla CSS properties for typography and grid systems without relying on external CSS frameworks.\n"
        "- View Layer (views.py): Implements business logic. It handles form evaluations, verifies roles, writes data, and handles file serving.\n"
        "- Authentication Backend (backends.py): Enables login operations using the email address instead of default usernames."
    )

    # Save document
    output_filename = "PANDA_NGAZI_DOCUMENTATION_CH_5.docx"
    doc.save(output_filename)
    print(f"Successfully generated {output_filename}!")

if __name__ == "__main__":
    generate_ch5_document()
